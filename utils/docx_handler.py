from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

def fill_report_template(
    template_path,
    output_path,
    week_ending,
    training_mode,
    daily_entries,
    details_notes,
    designation,
    signature_path=None  # path to PNG
):
    doc = Document(template_path)

    # --- Fill Table 0 (Main Diary Table)
    diary_table = doc.tables[0]
    # [0,0]: week ending
    orig_text = diary_table.cell(0,0).text
    if "\n" in orig_text:
        prefix = orig_text.split("\n")[0]
    else:
        prefix = orig_text
    diary_table.cell(0,0).text = f"{prefix}\nSunday: {week_ending}"

    # [0,3]: training mode
    diary_table.cell(0,3).text = f"TRAINING MODE\n{training_mode}"

    days = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY"]
    for i, day in enumerate(days):
        row = diary_table.rows[i+2]
        # [i+2, 1]: date
        row.cells[1].text = daily_entries.get(day, {}).get("date", "")
        # [i+2, 2]: description
        row.cells[2].text = daily_entries.get(day, {}).get("desc", "")

    # --- Fill Table 1 (Summary, Signature Table)
    details_table = doc.tables[1]
    # [1,0]: weekly summary/details
    details_table.cell(1,0).text = details_notes

    # [2,2]: signature image (optional)
    if signature_path:
        # Clear any existing text
        details_table.cell(2,2).text = ""
        # Insert image
        paragraph = details_table.cell(2,2).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(signature_path, width=Inches(1.2))

    # [5,0]: week-ending date (sunday)
    details_table.cell(5,0).text = f"DATE: {week_ending}"

    # [5,1]: designation
    details_table.cell(5,1).text = f"DESIGNATION AND SIGNATURE\n{designation}"
    # Optionally also add signature image here if required

    doc.save(output_path)
